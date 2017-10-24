# Use the SessionDetails.csv file to create a friendly formatted version to pass out to students.
import os.path
from docx import *

# #################### #
# User changeable items #
anyconnect_url = 'dcloud-sjc-anyconnect.cisco.com'
title = "SDA Test Drive"
# #################### #

inputfilename = 'SessionDetails.csv'
outputdirectory = 'ToPrint'


def build_document(doc, username, password, sessionline):
    # The text of the document
    doc.add_paragraph('{} event'.format(title), style='Title')
    doc.add_paragraph('Use the following instructions to connect to and use your assigned dCloud demonstration:')
    doc.add_paragraph('Open your AnyConnect VPN client.', style='List Bullet')
    p = doc.add_paragraph('Connect to: ', style='List Bullet')
    p.add_run(anyconnect_url).bold = True
    p = doc.add_paragraph('Username: ', style='List Bullet')
    p.add_run(username).bold = True
    p = doc.add_paragraph('Password: ', style='List Bullet')
    p.add_run(password).bold = True
    doc.add_paragraph('Follow along with the proctor, or use the provided guide,'
                      ' to go through the hands-on demonstration.', style='List Bullet')
    doc.add_paragraph('If you encounter issues please ask a proctor for assistance.')

    p = doc.add_paragraph()
    p.add_run('Notes: ').bold = True
    p = doc.add_paragraph('You can download the guide from: ', style='List Bullet')
    p.add_run('http://tinyurl.com/sda-testdrive').bold = True
    p = doc.add_paragraph('Your dCloud Session ID is: ', style='List Bullet')
    p.add_run(sessionline['sessionid']).bold = True
    p = doc.add_paragraph('Your session name is: ', style='List Bullet')
    p.add_run(sessionline['sessionname']).bold = True

    return doc


def main():
    fn = os.path.join(os.path.dirname(__file__), inputfilename)
    # Create and set up the document
    doc = Document()

    # Parse the SessionDetails.csv file, line by line.
    with open(fn, 'r') as f:
        for fline in f:
            fsplit = fline.split(',')  # There are 9 columns in the SessionDetails.csv
            sessionline = {
                'sessionid': fsplit[0],
                'sessionname': fsplit[1],
                'usernames': fsplit[2],
                'password': fsplit[3],
                'start': fsplit[4],
                'stop': fsplit[5],
                'dnsassets': fsplit[6],
                'sharedwith': fsplit[7],
                'endpointkit': fsplit[8]
            }

            # We are only interested in sessions that are started.
            if sessionline['usernames'] != 'Not Started' and sessionline['usernames'] != ' Usernames':
                vpn_user1 = sessionline['usernames'].split(';')[0]
                tmp = vpn_user1.split('user1')[0]
                tmp = tmp.split('v')[1]
                url = 'http://oob.vpod' + tmp + '.dc-01.com'
                sessionline['url'] = url
                password = sessionline['password'].split('"')[1]

                # Add page to document
                doc = build_document(doc, vpn_user1, password, sessionline)
                doc.add_page_break()

    # Save the document
    outputfilename = 'ToPrint.docx'
    gn = os.path.join(os.path.dirname(__file__), outputdirectory, outputfilename)
    doc.save(gn)
    print("Document saved: {}.".format(gn))


if __name__ == '__main__':
    print("Welcome to the dCloud session_parser program!")
    print("Be sure to update the variables: 'anyconnect_url' and 'title' to match your event's needs.")

    if not os.path.isfile(os.path.join(os.path.dirname(__file__), inputfilename)):
        print("Error: This program requires the '{}' file to be in the same directory"
              " as the '{}' program.").format(inputfilename, os.path.basename(__file__))
        exit(1)

    if not os.path.exists(os.path.join(os.path.dirname(__file__), outputdirectory)):
        os.makedirs(os.path.join(os.path.dirname(__file__), outputdirectory))

    print("")
    print("### Begin Program ###")

    main()

    print("### Program is finished ###")
    print("")
    print("Note: The files created by this program are be stored in the '{}' folder.".format(outputdirectory))
