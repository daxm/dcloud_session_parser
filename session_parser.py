# Use the SessionDetails.csv file to create a friendly formatted version to pass out to students.
import os.path
from docx import *

# #################### #
# User changeable items #
anyconnect_url = 'dcloud-sjc-anyconnect.cisco.com'
title = "SDA Test Drive"
# #################### #

inputfilename = 'SessionDetails.csv'
outputdirectory = 'ToPrint'


def build_document(username, password, sessionline):
    # Create and set up the document
    doc = Document()
    doc.add_paragraph('{} event'.format(title), style='Title')
    doc.add_paragraph('Use the following instructions to connect to your pod.')
    doc.add_paragraph('Open your AnyConnect VPN client.', style='List Number')
    p = doc.add_paragraph('Connect to: ', style='List Number')
    p.add_run(anyconnect_url).bold = True
    p = doc.add_paragraph('Username: ', style='List Number')
    p.add_run(username).bold = True
    p = doc.add_paragraph('Password: ', style='List Number')
    p.add_run(password).bold = True
    doc.add_paragraph('If you encounter issues please ask a proctor for assistance.')
    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('If asked, your dCloud Session ID is: ')
    p.add_run(sessionline['sessionid']).bold = True

    # Save the document
    outputfilename = '{} -- {}.docx'.format(sessionline['sessionname'], sessionline['sessionid'])
    gn = os.path.join(os.path.dirname(__file__), outputdirectory, outputfilename)
    doc.save(gn)
    print("Document saved: {}.".format(gn))


def main():
    fn = os.path.join(os.path.dirname(__file__), inputfilename)
    with open(fn, 'r') as f:
        for fline in f:
            fsplit = fline.split(',')  # There are 9 columns in the SessionDetails.csv
            sessionline = {
                'sessionid': fsplit[0],
                'sessionname': fsplit[1],
                'usernames': fsplit[2],
                'password': fsplit[3],
                'start': fsplit[4],
                'stop': fsplit[5],
                'dnsassets': fsplit[6],
                'sharedwith': fsplit[7],
                'endpointkit': fsplit[8]
            }
            # We are only interested in sessions that are started.
            if sessionline['usernames'] != 'Not Started' and sessionline['usernames'] != ' Usernames':
                vpn_user1 = sessionline['usernames'].split(';')[0]
                tmp = vpn_user1.split('user1')[0]
                tmp = tmp.split('v')[1]
                url = 'http://oob.vpod' + tmp + '.dc-01.com'
                sessionline['url'] = url
                password = sessionline['password'].split('"')[1]

                # Build and save the document
                build_document(vpn_user1, password, sessionline)


if __name__ == '__main__':
    print("Welcome to the dCloud session_parser program!")
    print("Be sure to update the variables: 'anyconnect_url' and 'title' to match your event's needs.")
    if os.path.isfile(os.path.join(os.path.dirname(__file__), inputfilename)) and os.path.exists(
            os.path.join(os.path.dirname(__file__), outputdirectory)):
        print("")
        print("### Begin Program ###")
        main()
        print("### Program is finished ###")
        print("")
        print("Note: The files created by this program are be stored in the '{}' folder.".format(outputdirectory))
    else:
        print("""
Error:  This program requires the '{}' to be in the same directory as the '{}' file.
This program also requires the '{}' folder in this same directory.
""".format(inputfilename, os.path.basename(__file__), outputdirectory))
